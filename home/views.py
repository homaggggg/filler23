from django.shortcuts import render
from .models import Doc
from django.http import StreamingHttpResponse
from wsgiref.util import FileWrapper
import os
import docx


def index(request):
    """
    Функция отображения для домашней страницы сайта.
    """
    num_docs=Doc.objects.all().count()
    return render(
        request,
        'index.html',
        context={'num_books':num_docs},
    )

def replace_keywords_in_docx(docx_path, keyword_dict, output_path):
    doc = docx.Document(docx_path)
    slovar = keyword_dict

    for paragraph in doc.paragraphs:
        for k, v in slovar.items():
            if k in paragraph.text:
                paragraph.text = paragraph.text.replace(k, v)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k, v in slovar.items():
                    if k in cell.text:
                        cell.text = cell.text.replace(k, v)
    doc.save(output_path)


def create(request):
    keyword_dict = {
    "{ИмяФамилияОтчество}": "Значение1",
    "{Группа}": "Значение2",
    "{Кафедра}": "Значение3",
    "{Группа}": "Значение4",
    "{Группа}": "Значение5",
    "{Группа}": "Значение6",
    }
    #файл с ключами
    input_file = "Приложение Д.docx"
    # название итогового файта
    output_file = "ITOG1.docx"
    replace_keywords_in_docx(input_file, keyword_dict, output_file)


# def download(request):
#     file = 'path_to_file'
#     filename = os.path.basename(file)
#     chunk_size = 8192
#     response = StreamingHttpResponse(FileWrapper(open(file, 'rb'), chunk_size),
#                         content_type=mimetypes.guess_type(file)[0])
#     response['Content-Length'] = os.path.getsize(file)    
#     response['Content-Disposition'] = "attachment; filename=%s" % filename
#     return response
